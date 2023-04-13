import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from django.http import FileResponse, HttpResponse
import mimetypes
import io
from docx2pdf import convert


class Report:
    _student = ''
    _teacher = ''
    _main_teacher = ''
    _assistant = ''
    _title = ''
    _group = ''
    _total_goals = []

    def __init__(self, student_name='', assistant_name='', teacher_name='',
                 main_teacher_name='', group_number='', title='', total_goals=None):
        if total_goals is None:
            total_goals = []
        self._student = student_name
        self._teacher = teacher_name
        self._assistant = assistant_name
        self._main_teacher = main_teacher_name
        self._group = group_number
        self._title = title
        self._total_goals = total_goals

    def generate_report(self):
        # Создание пустого документа
        document = Document()
        # Найстройка свойств документа по нормам НИТПУ
        # доступ к первой секции:
        section = document.sections[0]
        # левое поле в см
        section.left_margin = Cm(2)
        # правое поле в см
        section.right_margin = Cm(2)
        # верхнее поле в см
        section.top_margin = Cm(2)
        # нижнее поле в см
        section.bottom_margin = Cm(1.5)
        # Настройка стиля
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        # Вставка общей информации
        paragraph = document.add_paragraph('')
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Cm(11.5)
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.line_spacing = 1
        paragraph_format.space_after = Pt(0)
        teacher_position_permanent = "Руководитель ООП"
        # main_teacher_name = "Е. И. Громаков"
        paragraph.add_run('УТВЕРЖДАЮ\n{0}\n__________{1}\n«___»_________20___ г.'.format(teacher_position_permanent,
                                                                                         self._main_teacher))
        paragraph2 = document.add_paragraph('')
        paragraph2.add_run('\nИНДИВИДУАЛЬНОЕ ЗАДАНИЕ НА УИРС/НИРС').bold = True
        paragraph_format = paragraph2.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format.line_spacing = 1
        paragraph_format.space_after = Pt(0)

        title = self._title
        if title:
            paragraph30 = document.add_paragraph('')
            paragraph30.add_run('\n' + title.upper() + '\n').bold = True
            paragraph_format = paragraph30.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph_format.line_spacing = 1
            paragraph_format.space_after = Pt(0)
        paragraph3 = document.add_paragraph('')
        paragraph3.add_run('1. Перечень работ (заданий), подлежащих выполнению:').bold = True
        paragraph_format = paragraph3.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.line_spacing = 1
        paragraph_format.space_after = Pt(0)
        # Вставка таблиц
        total_goals = self._total_goals
        goals1_quantity = len(total_goals[0])
        goals2_quantity = len(total_goals[1])
        table = document.add_table(rows=goals1_quantity, cols=1)
        table.style = 'Table Grid'
        for goal in range(goals1_quantity):
            table.cell(goal, 0).text = str(goal+1) + '. ' + total_goals[0][goal]





        paragraph3 = document.add_paragraph('')

        paragraph3.add_run('\n2. Перечень отчетных материалов и требования к их оформлению:').bold = True
        paragraph_format = paragraph3.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.line_spacing = 1
        paragraph_format.space_after = Pt(0)

        table = document.add_table(rows=goals2_quantity, cols=1)
        table.style = 'Table Grid'
        for goal in range(goals2_quantity):
            table.cell(goal, 0).text = str(goal+1) + '. ' + total_goals[1][goal]

        # Подписи документа после таблиц
        teacher_position = "Ассистент ОАР ИШИТР"
        teacher_name = self._teacher
        paragraph4 = document.add_paragraph('')
        paragraph4.add_run('\nРуководитель УИРС/НИРС\n')
        paragraph4.add_run(teacher_position).underline = True
        paragraph4.add_run('                   _______________                              ')
        paragraph4.add_run(teacher_name).underline = True
        paragraph_format = paragraph4.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.line_spacing = 1
        paragraph_format.space_after = Pt(0)
        font.size = Pt(12)
        paragraph5 = document.add_paragraph('')
        paragraph_format = paragraph5.paragraph_format
        line_spacing = Pt(3)
        paragraph5.add_run(	'                     (должность)                                   ' +
                               '                           (подпись)				         (Ф. И. О.) ').font.superscript = True
        paragraph_format.line_spacing = 1
        paragraph_format.space_after = Pt(0)

        student_group = self._group
        student_name = self._student
        assistant_name = self._assistant
        if len(assistant_name) > 1:
            print('ENTRANCE')
            paragraph7 = document.add_paragraph('')
            paragraph7.add_run('\nКонсультант\n')
            paragraph7.add_run(teacher_position).underline = True
            paragraph7.add_run('                   _______________                              ')
            paragraph7.add_run(assistant_name).underline = True
            paragraph_format = paragraph7.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph_format.line_spacing = 1
            paragraph_format.space_after = Pt(0)
            font.size = Pt(12)
            paragraph8 = document.add_paragraph('')
            paragraph_format = paragraph8.paragraph_format
            line_spacing = Pt(3)
            paragraph8.add_run('                     (должность)                                   ' +
                               '                           (подпись)				         (Ф. И. О.) ').font.superscript = True
            paragraph_format.line_spacing = 1
            paragraph_format.space_after = Pt(0)

        paragraph6 = document.add_paragraph('')
        paragraph6.add_run('\nЗадание принял к исполнению         _______________                              ')
        paragraph6.add_run(student_name).underline = True
        paragraph_format = paragraph6.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.line_spacing = 1
        paragraph_format.space_after = Pt(0)
        paragraph6 = document.add_paragraph('Студент группы ' + student_group + '                            ')
        paragraph6.add_run('      (подпись)                                                              ' +
                           '(Ф. И. О. обучающегося)\n').font.superscript = True
        paragraph6.add_run('\n«___» _________ 20___г.')
        paragraph_format = paragraph6.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format.line_spacing = 1
        paragraph_format.space_after = Pt(0)
        # Сохранение документа и конвертация в pdf
        document.save('server/static/report1.docx')


